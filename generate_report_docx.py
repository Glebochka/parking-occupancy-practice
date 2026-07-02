from __future__ import annotations

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
WORKSPACE = ROOT.parent
REPORT_PATH = WORKSPACE / "Отчет.docx"
BACKUP_PATH = WORKSPACE / "Отчет_backup_before_autofill.docx"
ARTIFACTS_DIR = ROOT / "artifacts"
ASSETS_DIR = ROOT / "report_assets"
SUMMARY_PATH = ARTIFACTS_DIR / "summary.csv"
MANIFEST_PATH = ROOT / "data" / "processed" / "manifest.csv"
RESULTS_SUMMARY_PATH = ROOT / "results_summary.md"

FONT_NAME = "Times New Roman"
FONT_SIZE = 13
BLACK = RGBColor(0, 0, 0)
ACCESS_DATE = "02.07.2026"


def set_run_style(run, bold: bool = False, italic: bool = False, size: int = FONT_SIZE) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)


def set_paragraph_format(paragraph, first_line_indent: bool = True, centered: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line_indent:
        fmt.first_line_indent = Cm(1.25)
    else:
        fmt.first_line_indent = Cm(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY


def add_paragraph_text(doc: Document, text: str, first_line_indent: bool = True) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, first_line_indent=first_line_indent)
    run = paragraph.add_run(text)
    set_run_style(run)


def add_heading_text(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, first_line_indent=False)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_style(run, bold=True)
    if level == 1:
        paragraph.style = doc.styles["Heading 1"]
    elif level == 2:
        paragraph.style = doc.styles["Heading 2"]


def add_caption(doc: Document, number: int, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, first_line_indent=False, centered=True)
    run1 = paragraph.add_run(f"Рис. {number}. ")
    set_run_style(run1, bold=True)
    run2 = paragraph.add_run(text)
    set_run_style(run2)


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_toc(paragraph) -> None:
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Обновите поле содержания в Word: выделите содержание и нажмите F9."
    r.append(t)
    fld_simple.append(r)
    paragraph._element.append(fld_simple)


def add_page_number(paragraph) -> None:
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld_simple.append(r)
    paragraph._element.append(fld_simple)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(1.5)
    section.different_first_page_header_footer = True

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(p)
    for run in p.runs:
        set_run_style(run)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE)
    normal.font.color.rgb = BLACK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style.font.size = Pt(FONT_SIZE)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_title_page(doc: Document) -> None:
    lines = [
        "Министерство цифрового развития, связи и массовых коммуникаций Российской Федерации",
        "Ордена Трудового Красного Знамени федеральное государственное бюджетное образовательное учреждение высшего образования",
        "МОСКОВСКИЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ",
        "СВЯЗИ И ИНФОРМАТИКИ",
        'Кафедра "Программная инженерия"',
    ]

    for line in lines:
        paragraph = doc.add_paragraph()
        set_paragraph_format(paragraph, first_line_indent=False, centered=True)
        run = paragraph.add_run(line)
        set_run_style(run)

    for _ in range(6):
        doc.add_paragraph()

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, first_line_indent=False, centered=True)
    run = paragraph.add_run("ОТЧЁТ")
    set_run_style(run, bold=True)

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, first_line_indent=False, centered=True)
    run = paragraph.add_run("по производственной практике (проектно-технологическая)")
    set_run_style(run)

    for _ in range(7):
        doc.add_paragraph()

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("Выполнил:\nСтудент группы: УБВТ2304\nФИО: Решетников Г.Е.")
    set_run_style(run)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run('Проверил:\nБорзов В.М., ассистент кафедры "Программная инженерия"')
    set_run_style(run)

    for _ in range(8):
        doc.add_paragraph()

    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, first_line_indent=False, centered=True)
    run = paragraph.add_run("Москва, 2026")
    set_run_style(run)


def add_table_from_dataframe(doc: Document, df: pd.DataFrame, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    set_run_style(run)


def add_figure(doc: Document, path: Path, caption_number: int, caption_text: str, width_cm: float = 14.0) -> int:
    doc.add_picture(str(path), width=Cm(width_cm))
    paragraph = doc.paragraphs[-1]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption_number, caption_text)
    return caption_number + 1


def format_metric(value: float) -> str:
    return f"{value:.4f}"


def build_report() -> None:
    if REPORT_PATH.exists():
        BACKUP_PATH.write_bytes(REPORT_PATH.read_bytes())

    summary_df = pd.read_csv(SUMMARY_PATH).sort_values("f1", ascending=False).reset_index(drop=True)
    manifest_df = pd.read_csv(MANIFEST_PATH)
    best = summary_df.iloc[0]

    split_summary = (
        manifest_df.groupby(["split", "label"])
        .size()
        .reset_index(name="count")
        .sort_values(["split", "label"])
    )

    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    add_page_break(doc)
    add_heading_text(doc, "СОДЕРЖАНИЕ", level=1)
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=False)
    add_toc(p)

    add_page_break(doc)
    add_heading_text(doc, "Индивидуальное задание", level=1)
    add_paragraph_text(
        doc,
        "Индивидуальное задание по производственной практике заключалось в разработке системы определения "
        "занятости парковочного места по изображению с применением методов компьютерного зрения и искусственных "
        "нейронных сетей. В рамках задания требовалось подготовить данные, обучить и сравнить не менее пяти "
        "архитектур нейронных сетей, провести анализ качества моделей, выбрать лучшую архитектуру и реализовать "
        "демонстрационный программный модуль с выводом результата классификации, вероятности предсказания и "
        "сохранением истории запусков."
    )

    add_page_break(doc)
    add_heading_text(doc, "Введение", level=1)
    intro_paragraphs = [
        "Задача автоматического определения занятости парковочного места относится к числу актуальных задач "
        "интеллектуальных транспортных систем. Рост автомобилизации, ограниченность городской парковочной "
        "инфраструктуры и необходимость оперативного предоставления информации водителям требуют применения "
        "средств автоматического мониторинга парковочных зон [1-5].",
        "Традиционные решения часто основаны на использовании отдельных датчиков для каждого парковочного места, "
        "что увеличивает стоимость внедрения и эксплуатации системы. Альтернативным подходом является обработка "
        "изображений с камер наблюдения, при которой состояние парковочного места определяется по визуальным "
        "признакам. Такой подход позволяет использовать одну камеру для контроля множества мест, упрощает "
        "масштабирование системы и снижает затраты на инфраструктуру [2-5].",
        "Целью выполненной работы являлась разработка программного решения для определения занятости парковочного "
        "места по изображению, а также экспериментальное сравнение нескольких архитектур нейронных сетей на "
        "реальном наборе данных. В ходе практики была реализована подготовка датасета PKLot [6], настроен цикл "
        "обучения и тестирования моделей в среде Python 3.11 [7] с использованием библиотек PyTorch [8], "
        "torchvision [9], scikit-learn [10], pandas [11] и Pillow [12].",
        "Отчёт отражает фактически выполненный этап практической работы: подготовку сбалансированной выборки, "
        "реализацию пайплайна обучения, сравнительный запуск пяти моделей, анализ полученных результатов и "
        "формирование материалов для дальнейшего использования в итоговой версии отчёта."
    ]
    for text in intro_paragraphs:
        add_paragraph_text(doc, text)

    add_page_break(doc)
    add_heading_text(doc, "1. Определение объекта, предмета и цели исследования", level=1)
    section1 = [
        "Объектом исследования является процесс автоматизированного мониторинга состояния парковочных мест по "
        "изображениям, получаемым с камер наблюдения.",
        "Предметом исследования являются методы и модели глубокого обучения для бинарной классификации состояния "
        'парковочного места по классам "свободно" и "занято".',
        "Цель практической разработки состоит в создании программного решения, позволяющего по изображению "
        "отдельного парковочного места определять его состояние, оценивать уверенность модели и сохранять "
        "результаты обработки для последующего анализа.",
        "Для достижения поставленной цели были сформулированы следующие задачи: изучить современные подходы к "
        "распознаванию занятости парковочных мест [1-5]; выбрать подходящий набор данных и подготовить обучающую "
        "выборку [6]; реализовать программный конвейер предобработки и обучения на Python [7-12]; провести "
        "сравнение пяти архитектур нейронных сетей; выбрать лучшую модель по совокупности метрик качества и "
        "скорости работы; подготовить графические материалы и фактическую основу для итоговой отчётности.",
        "Практическая значимость работы определяется возможностью интеграции полученного решения в системы "
        "интеллектуальной парковки, аналитики транспортной инфраструктуры и сервисы информирования водителей о "
        "наличии свободных мест."
    ]
    for text in section1:
        add_paragraph_text(doc, text)

    add_page_break(doc)
    add_heading_text(doc, "2. Постановка задач практической разработки", level=1)
    section2_intro = [
        "В качестве исходных данных был использован публичный набор PKLot, содержащий изображения парковок при "
        "различных погодных условиях и XML-разметку отдельных парковочных мест [6]. Для текущего этапа практики "
        "был реализован собственный сценарий подготовки данных, автоматически вырезающий отдельные парковочные "
        "места из кадров по полигональной разметке и формирующий структуру каталогов, пригодную для обучения "
        "моделей классификации.",
        "С технической точки зрения разработка включала несколько взаимосвязанных задач. Во-первых, потребовалось "
        "ускорить чтение исходного архива PKLot.tar.gz: последовательная обработка архива оказалась существенно "
        "эффективнее случайного доступа к сжатому файлу, поскольку исключает многократную повторную распаковку "
        "одних и тех же фрагментов данных. Во-вторых, необходимо было реализовать балансировку классов и "
        "формирование разбиения train/validation/test. В-третьих, следовало построить единый цикл обучения и "
        "оценивания для нескольких архитектур с одинаковыми параметрами запуска.",
        "В рамках практики были реализованы следующие программные модули: сценарий подготовки датасета "
        "prepare_pklot.py, модуль обучения и сравнения benchmark.py, модуль инференса predict.py, а также "
        "вспомогательные компоненты для построения моделей, вычисления метрик и генерации материалов для отчёта. "
        "Демонстрационный модуль выполнен в консольном виде: по входному изображению система возвращает класс, "
        "вероятность, а также записывает историю предсказаний в JSONL-файл.",
        "В качестве сравниваемых архитектур были выбраны ResNet18, MobileNetV3 Small, EfficientNet-B0, "
        "DenseNet121 и ConvNeXt-Tiny. Такой выбор позволяет сопоставить как относительно лёгкие модели, пригодные "
        "для быстрого инференса, так и более тяжёлые архитектуры с высокой выразительной способностью [2-5, 9]."
    ]
    for text in section2_intro:
        add_paragraph_text(doc, text)

    add_heading_text(doc, "2.1. Использованные данные и инструменты", level=2)
    add_paragraph_text(
        doc,
        "Для обучения была сформирована сбалансированная подвыборка объёмом 10 000 изображений: по 5 000 "
        'примеров классов "empty" и "occupied". Разбиение на обучающую, валидационную и тестовую выборки было '
        "выполнено стратифицированно, что позволило сохранить равное распределение классов во всех подмножествах."
    )

    split_rows = [
        [row["split"], row["label"], str(int(row["count"]))]
        for _, row in split_summary.iterrows()
    ]
    add_table_from_dataframe(doc, split_summary, ["Подмножество", "Класс", "Количество"], split_rows)

    figure_counter = 1
    figure_counter = add_figure(
        doc,
        ASSETS_DIR / "dataset_distribution.png",
        figure_counter,
        "Распределение подготовленных изображений по подмножествам и классам.",
    )
    figure_counter = add_figure(
        doc,
        ASSETS_DIR / "dataset_samples.png",
        figure_counter,
        "Примеры изображений парковочных мест после автоматической сегментации.",
    )

    add_paragraph_text(
        doc,
        "Программная реализация выполнена на языке Python 3.11 [7]. Для обучения нейронных сетей использовалась "
        "библиотека PyTorch [8], готовые предобученные архитектуры были загружены средствами torchvision [9], "
        "расчёт метрик качества выполнялся с помощью scikit-learn [10], хранение и обработка табличных результатов "
        "осуществлялись в pandas [11], а операции чтения, обрезки и сохранения изображений — библиотекой Pillow [12]."
    )

    add_heading_text(doc, "2.2. Реализация программного решения", level=2)
    impl_paragraphs = [
        "Сценарий prepare_pklot.py выполняет два последовательных прохода по архиву PKLot. На первом проходе "
        "из XML-файлов считываются координаты парковочных мест и с помощью алгоритма резервуарной выборки "
        "отбирается сбалансированный набор примеров двух классов. На втором проходе изображения считываются уже "
        "в порядке их расположения в архиве, после чего для каждого выбранного места формируется отдельный кроп. "
        "Такое решение позволило кратно снизить время подготовки данных по сравнению со схемой случайного доступа "
        "к сжатому архиву.",
        "Сценарий benchmark.py запускает единый цикл обучения для пяти выбранных архитектур. В качестве входного "
        "размера использовались изображения 224x224 пикселя, оптимизация выполнялась алгоритмом AdamW, а для "
        "обучения применялись стандартные аугментации: изменение яркости и контраста, горизонтальное отражение и "
        "нормализация по статистикам ImageNet.",
        "Сценарий predict.py реализует демонстрационный режим использования лучшей модели. На вход подаётся путь "
        "к изображению, на выходе формируется предсказанный класс, значение уверенности модели и JSONL-запись в "
        "файле истории запусков. Таким образом, по факту выполненной работы реализован не только этап исследования, "
        "но и рабочий прототип модуля инференса."
    ]
    for text in impl_paragraphs:
        add_paragraph_text(doc, text)

    add_page_break(doc)
    add_heading_text(doc, "3. Описание результатов практической деятельности", level=1)
    section3_intro = [
        "После подготовки данных был выполнен сравнительный запуск пяти архитектур нейронных сетей. На текущем "
        "этапе практики использовался ускоренный режим обучения: по 2 эпохи на каждую модель. Несмотря на "
        "сокращённое число эпох, такой режим позволил получить фактическую таблицу сравнения и выбрать архитектуру, "
        "демонстрирующую наилучший баланс между качеством и скоростью.",
        "Оценка выполнялась по метрикам accuracy, precision, recall и F1-score, а также по среднему времени "
        "инференса на одно изображение и размеру итогового файла модели. Основным критерием выбора лучшей модели "
        "на данном этапе был F1-score, поскольку он учитывает одновременно полноту и точность классификации [10]."
    ]
    for text in section3_intro:
        add_paragraph_text(doc, text)

    result_rows = []
    for _, row in summary_df.iterrows():
        result_rows.append(
            [
                row["model_name"],
                format_metric(row["accuracy"]),
                format_metric(row["precision"]),
                format_metric(row["recall"]),
                format_metric(row["f1"]),
                f"{row['latency_ms_per_image']:.2f}",
                f"{row['model_size_mb']:.2f}",
            ]
        )
    add_table_from_dataframe(
        doc,
        summary_df,
        ["Модель", "Accuracy", "Precision", "Recall", "F1", "мс/изобр.", "МБ"],
        result_rows,
    )

    figure_counter = add_figure(
        doc,
        ASSETS_DIR / "model_comparison.png",
        figure_counter,
        "Сравнение моделей по F1-score и средней задержке инференса.",
    )

    add_paragraph_text(
        doc,
        f"По итогам сравнения наилучший результат показала модель {best['model_name']}, достигшая accuracy "
        f"{best['accuracy']:.4f}, precision {best['precision']:.4f}, recall {best['recall']:.4f} и F1-score "
        f"{best['f1']:.4f}. Среднее время обработки одного изображения составило {best['latency_ms_per_image']:.2f} мс, "
        f"а размер файла модели — {best['model_size_mb']:.2f} МБ. Данный результат позволяет рассматривать "
        f"{best['model_name']} как основную архитектуру для дальнейшего развития проекта."
    )

    figure_counter = add_figure(
        doc,
        ASSETS_DIR / "best_model_confusion_matrix.png",
        figure_counter,
        f"Матрица ошибок лучшей модели {best['model_name']}.",
        width_cm=10.0,
    )

    add_paragraph_text(
        doc,
        "Матрица ошибок показывает, что основная доля ошибок связана с ложным определением занятых мест как "
        "свободных или свободных как занятых в условиях схожих текстур, частичных перекрытий, теней и различий "
        "освещения. При этом доля корректных ответов остаётся высокой, что подтверждает пригодность выбранного "
        "подхода для решения поставленной задачи."
    )

    figure_counter = add_figure(
        doc,
        ASSETS_DIR / "prediction_examples.png",
        figure_counter,
        "Пример предсказаний лучшей модели на тестовых изображениях.",
    )

    add_paragraph_text(
        doc,
        "Дополнительно была проверена работа демонстрационного модуля predict.py на реальном тестовом изображении. "
        "Модель EfficientNet-B0 корректно определила занятое парковочное место и сохранила результат вместе с "
        "вероятностями классов в файл истории предсказаний. Это подтверждает, что реализованный программный "
        "конвейер функционирует не только в режиме офлайн-обучения, но и в прикладном режиме инференса."
    )

    add_paragraph_text(
        doc,
        "Таким образом, по факту выполненной работы реализованы основные элементы практической части: подготовка "
        "данных, ускоренный модуль предобработки, обучение и сравнение пяти архитектур, выбор лучшей модели, "
        "сохранение метрик и подготовка графических материалов для отчётности. Следующим шагом развития проекта "
        "может стать увеличение числа эпох обучения, расширение набора сценариев валидации и разработка "
        "графического пользовательского интерфейса."
    )

    add_page_break(doc)
    add_heading_text(doc, "Заключение", level=1)
    conclusion_paragraphs = [
        "В ходе производственной практики была выполнена практическая разработка системы определения занятости "
        "парковочного места по изображению. Для решения задачи был выбран набор данных PKLot, реализована "
        "автоматическая подготовка сбалансированной выборки, построен программный конвейер обучения и сравнения "
        "нескольких архитектур нейронных сетей.",
        "По итогам текущего этапа экспериментов были сравнены пять моделей: ResNet18, MobileNetV3 Small, "
        "EfficientNet-B0, DenseNet121 и ConvNeXt-Tiny. Наилучшие результаты показала модель EfficientNet-B0, "
        "что позволило выбрать её в качестве базовой архитектуры для дальнейшего использования.",
        "Практический результат работы заключается в наличии готового программного проекта, включающего сценарий "
        "подготовки данных, модуль обучения, модуль инференса, сохранение метрик и визуальные материалы для отчёта. "
        "Полученные результаты подтверждают применимость методов глубокого обучения для задачи определения занятости "
        "парковочного места и создают основу для дальнейшего улучшения системы."
    ]
    for text in conclusion_paragraphs:
        add_paragraph_text(doc, text)

    add_page_break(doc)
    add_heading_text(doc, "Список использованных источников", level=1)
    sources = [
        "Marek M. Image-Based Parking Space Occupancy Classification: Dataset and Baseline [Электронный ресурс]. "
        "2021. URL: https://arxiv.org/abs/2107.12207 (дата обращения: 02.07.2026).",
        "Martynova A., Kuznetsov M., Porvatov V., Tishin V., Kuznetsov A., Semenova N., Kuznetsova K. "
        "Revising deep learning methods in parking lot occupancy detection [Электронный ресурс]. 2023. "
        "URL: https://arxiv.org/abs/2306.04288 (дата обращения: 02.07.2026).",
        "Yuldashev Y., Mukhiddinov M., Abdusalomov A.B., Nasimov R., Cho J. Parking Lot Occupancy Detection with "
        "Improved MobileNetV3 // Sensors. 2023. Vol. 23. No. 17. Art. 7642. URL: https://www.mdpi.com/1424-8220/23/17/7642 "
        "(дата обращения: 02.07.2026).",
        "Grbić R., Koch B. Automatic Vision-Based Parking Slot Detection and Occupancy Classification "
        "[Электронный ресурс]. 2023. URL: https://arxiv.org/abs/2308.08192 (дата обращения: 02.07.2026).",
        "da Luz G.P.C.P., Sato G.M., Gonzalez L.F.G., Borin J.F. Smart Parking with Pixel-Wise ROI Selection for "
        "Vehicle Detection Using YOLOv8, YOLOv9, YOLOv10, and YOLOv11 [Электронный ресурс]. 2024. "
        "URL: https://arxiv.org/abs/2412.01983 (дата обращения: 02.07.2026).",
        "Parking Lot Database [Электронный ресурс]. Vision, Robotics and Imaging Laboratory, UFPR. "
        "URL: https://web.inf.ufpr.br/vri/databases/parking-lot-database/ (дата обращения: 02.07.2026).",
        "Python 3 Documentation [Электронный ресурс]. URL: https://docs.python.org/ (дата обращения: 02.07.2026).",
        "PyTorch Documentation [Электронный ресурс]. URL: https://docs.pytorch.org/docs/stable/index.html "
        "(дата обращения: 02.07.2026).",
        "Torchvision Models and Pre-trained Weights [Электронный ресурс]. URL: https://docs.pytorch.org/vision/main/models.html "
        "(дата обращения: 02.07.2026).",
        "Scikit-learn. Metrics and scoring: quantifying the quality of predictions [Электронный ресурс]. "
        "URL: https://scikit-learn.org/stable/modules/model_evaluation.html (дата обращения: 02.07.2026).",
        "pandas Documentation [Электронный ресурс]. URL: https://pandas.pydata.org/docs/ (дата обращения: 02.07.2026).",
        "Pillow Documentation [Электронный ресурс]. URL: https://pillow.readthedocs.io/ (дата обращения: 02.07.2026).",
    ]
    for idx, source in enumerate(sources, start=1):
        paragraph = doc.add_paragraph()
        set_paragraph_format(paragraph, first_line_indent=False)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(f"{idx}. {source}")
        set_run_style(run)

    add_page_break(doc)
    add_heading_text(doc, "Приложения", level=1)
    appendix_paragraphs = [
        f"Приложение А. Локальный каталог итогового проекта: {ROOT}",
        "Приложение Б. Основные файлы проекта: prepare_pklot.py, benchmark.py, predict.py, "
        "generate_report_assets.py, src/parking_occupancy/models.py, src/parking_occupancy/trainer.py.",
        "Приложение В. Сформированные материалы для отчёта: dataset_distribution.png, dataset_samples.png, "
        "model_comparison.png, best_model_confusion_matrix.png, prediction_examples.png.",
        "Примечание: на момент подготовки данной версии отчёта проект размещён в локальном рабочем каталоге. "
        "При публикации проекта на GitHub в приложение может быть добавлена постоянная ссылка на репозиторий."
    ]
    for text in appendix_paragraphs:
        add_paragraph_text(doc, text, first_line_indent=False)

    doc.save(REPORT_PATH)


if __name__ == "__main__":
    build_report()
