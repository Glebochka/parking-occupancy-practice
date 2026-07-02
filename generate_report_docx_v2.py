from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
WORKSPACE = ROOT.parent
REPORT_PATH = WORKSPACE / "Отчет.docx"
BACKUP_PATH = WORKSPACE / "Отчет_backup_before_autofill_v2.docx"
ASSETS_DIR = ROOT / "report_assets"
ARTIFACTS_DIR = ROOT / "artifacts"
MANIFEST_PATH = ROOT / "data" / "processed" / "manifest.csv"
SUMMARY_PATH = ARTIFACTS_DIR / "summary.csv"
WEATHER_PATH = ARTIFACTS_DIR / "weather_metrics.csv"
SOURCE_PATH = ARTIFACTS_DIR / "source_metrics.csv"
WORKBOOK_PATH = ARTIFACTS_DIR / "parking_occupancy_results.xlsx"
REPO_URL_PATH = ROOT / "repo_url.txt"

FONT_NAME = "Times New Roman"
FONT_SIZE = 13
BLACK = RGBColor(0, 0, 0)


def repo_url() -> str:
    if REPO_URL_PATH.exists():
        return REPO_URL_PATH.read_text(encoding="utf-8").strip()
    return "Ссылка на GitHub будет добавлена после публикации репозитория."


def set_run_style(run, bold: bool = False, italic: bool = False, size: int = FONT_SIZE) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)


def set_paragraph(paragraph, centered: bool = False, first_line_indent: bool = True) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Cm(1.25) if first_line_indent else Cm(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY


def add_text(doc: Document, text: str, centered: bool = False, first_line_indent: bool = True, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, centered=centered, first_line_indent=first_line_indent)
    run = paragraph.add_run(text)
    set_run_style(run, bold=bold)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, centered=False, first_line_indent=False)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_style(run, bold=True)
    paragraph.style = doc.styles["Heading 1" if level == 1 else "Heading 2"]


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_toc(paragraph) -> None:
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Обновите содержание в Microsoft Word клавишей F9."
    r.append(t)
    fld_simple.append(r)
    paragraph._element.append(fld_simple)


def add_page_number(paragraph) -> None:
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld_simple.append(r)
    paragraph._element.append(fld_simple)


def add_caption(doc: Document, number: int, text: str) -> int:
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, centered=True, first_line_indent=False)
    run1 = paragraph.add_run(f"Рис. {number}. ")
    set_run_style(run1, bold=True)
    run2 = paragraph.add_run(text)
    set_run_style(run2)
    return number + 1


def add_picture(doc: Document, image_path: Path, width_cm: float, caption_number: int, caption_text: str) -> int:
    doc.add_picture(str(image_path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return add_caption(doc, caption_number, caption_text)


def format_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(1.5)
    section.different_first_page_header_footer = True

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer_paragraph)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE)
    normal.font.color.rgb = BLACK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    for style_name in ("Heading 1", "Heading 2"):
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style.font.size = Pt(FONT_SIZE)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_title_page(doc: Document) -> None:
    title_lines = [
        "Министерство цифрового развития, связи и массовых коммуникаций Российской Федерации",
        "Ордена Трудового Красного Знамени федеральное государственное бюджетное образовательное учреждение высшего образования",
        "МОСКОВСКИЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ",
        "СВЯЗИ И ИНФОРМАТИКИ",
        'Кафедра "Программная инженерия"',
    ]
    for line in title_lines:
        add_text(doc, line, centered=True, first_line_indent=False)

    for _ in range(6):
        doc.add_paragraph()

    add_text(doc, "ОТЧЁТ", centered=True, first_line_indent=False, bold=True)
    add_text(doc, "по производственной практике (проектно-технологическая)", centered=True, first_line_indent=False)

    for _ in range(7):
        doc.add_paragraph()

    add_text(doc, "Выполнил:\nСтудент группы: УБВТ2304\nФИО: Решетников Г.Е.", first_line_indent=False)
    add_text(doc, 'Проверил:\nБорзов В.М., ассистент кафедры "Программная инженерия"', first_line_indent=False)

    for _ in range(8):
        doc.add_paragraph()

    add_text(doc, "Москва, 2026", centered=True, first_line_indent=False)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_style(run)


def build_report() -> None:
    if REPORT_PATH.exists():
        BACKUP_PATH.write_bytes(REPORT_PATH.read_bytes())

    manifest_df = pd.read_csv(MANIFEST_PATH)
    summary_df = pd.read_csv(SUMMARY_PATH).sort_values("f1", ascending=False).reset_index(drop=True)
    weather_df = pd.read_csv(WEATHER_PATH).sort_values("weather").reset_index(drop=True)
    source_df = pd.read_csv(SOURCE_PATH).sort_values("source").reset_index(drop=True)
    best = summary_df.iloc[0]

    split_df = (
        manifest_df.groupby(["split", "label"]).size().reset_index(name="count").sort_values(["split", "label"])
    )

    training_rows = [
        ["Эпохи", "2", "Ускоренный экспериментальный режим"],
        ["Batch size", "32", "Одинаковый для всех моделей"],
        ["Размер изображения", "224x224", "Стандартный вход для ImageNet-pretrained моделей"],
        ["Оптимизатор", "AdamW", "Использовался для всех запусков"],
        ["Learning rate", "3e-4", "Начальное значение"],
        ["Weight decay", "1e-4", "Регуляризация"],
        ["Scheduler", "CosineAnnealingLR", "Плавное уменьшение шага обучения"],
        ["Аугментации", "Flip + ColorJitter", "Повышение устойчивости к вариативности данных"],
    ]

    doc = Document()
    format_doc(doc)
    add_title_page(doc)

    add_page_break(doc)
    add_heading(doc, "Содержание", level=1)
    toc_paragraph = doc.add_paragraph()
    set_paragraph(toc_paragraph, centered=False, first_line_indent=False)
    add_toc(toc_paragraph)

    add_page_break(doc)
    add_heading(doc, "Индивидуальное задание", level=1)
    add_text(
        doc,
        "Индивидуальное задание по производственной практике состояло в разработке системы определения "
        "занятости парковочного места по изображению с применением методов компьютерного зрения и искусственных "
        "нейронных сетей. В рамках выполнения практики требовалось подготовить данные, обучить и сравнить не "
        "менее пяти архитектур нейронных сетей, реализовать демонстрационный модуль для инференса, обеспечить "
        "сохранение истории запусков и сформировать итоговые материалы для отчётности.",
    )

    add_page_break(doc)
    add_heading(doc, "Введение", level=1)
    intro_blocks = [
        "В современных городах проблема эффективного использования парковочного пространства является одной из "
        "ключевых задач интеллектуальных транспортных систем. Рост числа автомобилей приводит к увеличению "
        "нагрузки на парковочную инфраструктуру, а поиск свободного парковочного места становится фактором, "
        "влияющим на дорожные заторы, расход топлива и общее удобство городской мобильности [1-6].",
        "Классические системы контроля занятости парковочных мест строятся на отдельных датчиках, размещаемых "
        "непосредственно в зоне каждого парковочного места. Такой подход обеспечивает достаточную точность, "
        "однако требует значительных затрат на установку и обслуживание. Камерно-ориентированные решения "
        "являются более гибкой альтернативой, поскольку позволяют использовать существующую инфраструктуру "
        "видеонаблюдения и одновременно охватывать большое число парковочных мест [2-6].",
        "Задача определения занятости парковочного места по изображению сводится к бинарной классификации: "
        'необходимо определить, относится ли отдельный фрагмент изображения к классу "empty" или "occupied". '
        "Для решения такой задачи широко применяются сверточные нейронные сети и современные архитектуры, "
        "получившие предварительное обучение на крупных наборах данных [2-5, 8-12].",
        "Выполненная работа ориентирована на практическую реализацию полного цикла подготовки системы: от "
        "получения и обработки датасета до сравнительного анализа пяти архитектур, выбора лучшей модели, "
        "реализации демонстрационного интерфейса и экспорта результатов в удобные форматы для дальнейшей "
        "интерпретации.",
        "Для практической части был выбран датасет PKLot, который содержит изображения парковок при различных "
        "погодных условиях и XML-разметку парковочных мест [6]. На его основе сформирована сбалансированная "
        "выборка из 10 000 сегментированных изображений, пригодная для ускоренного экспериментального цикла. "
        "Полученные результаты легли в основу настоящего отчёта.",
    ]
    for block in intro_blocks:
        add_text(doc, block)

    add_page_break(doc)
    add_heading(doc, "1. Определение объекта, предмета и цели исследования", level=1)
    section1 = [
        "Объектом исследования является процесс автоматизированного мониторинга состояния парковочных мест по "
        "изображениям, получаемым с камер наблюдения.",
        "Предмет исследования составляют методы построения и сравнительного анализа нейронных сетей для "
        'бинарной классификации состояния парковочного места по классам "свободно" и "занято".',
        "Цель исследования заключается в разработке работоспособной программной системы, способной принимать "
        "изображение парковочного места, определять его состояние, выводить уверенность предсказания, накапливать "
        "историю проверок и формировать набор отчётных артефактов на основе результатов вычислительного эксперимента.",
        "Для достижения поставленной цели решались следующие задачи: анализ существующих публикаций по тематике "
        "распознавания занятости парковочных мест [1-5]; выбор и описание датасета PKLot [6]; разработка ускоренного "
        "конвейера подготовки изображений; построение программной среды на Python 3.11 [7]; использование "
        "PyTorch и torchvision для дообучения моделей [8, 9]; вычисление метрик посредством scikit-learn [10]; "
        "сохранение и агрегация результатов в pandas и Excel [11]; подготовка изображений и графических материалов "
        "при помощи Pillow и matplotlib [12].",
        "Практическая значимость выполненной работы состоит в возможности дальнейшей интеграции разработанного "
        "решения в системы интеллектуальных парковок, сервисы навигации по парковочным зонам, а также в проекты "
        "видеоаналитики для транспортной инфраструктуры. Дополнительной значимостью является воспроизводимость "
        "полученного результата благодаря наличию программного проекта и подготовленного набора отчётных файлов.",
    ]
    for block in section1:
        add_text(doc, block)

    add_heading(doc, "1.1. Степень разработанности проблемы", level=2)
    state_blocks = [
        "Анализ современных исследований показывает, что задача определения занятости парковочных мест активно "
        "развивается в нескольких направлениях. Одно направление связано с прямой классификацией заранее "
        "размеченных областей интереса, соответствующих отдельным парковочным местам [1-3]. Другое направление "
        "строится вокруг автоматического обнаружения парковочных мест и последующей классификации их состояния [4, 5].",
        "Исследования последних лет подчеркивают важность не только итоговой точности модели, но и устойчивости "
        "к изменению погодных условий, ракурса и освещения. В работах 2023-2024 годов отмечается, что архитектуры "
        "семейства EfficientNet и MobileNet обеспечивают сильный баланс между качеством и вычислительной "
        "эффективностью при задачах парковочной аналитики [2, 3].",
        "Таким образом, выбранная тема является практически значимой и исследовательски актуальной. Вместе с тем "
        "она оставляет пространство для собственной проектной реализации: адаптации структуры данных, проектирования "
        "конвейера подготовки изображений, выбора набора архитектур и анализа результатов на ограниченном, но "
        "репрезентативном подмножестве реальных данных.",
    ]
    for block in state_blocks:
        add_text(doc, block)

    add_page_break(doc)
    add_heading(doc, "2. Постановка задач практической разработки", level=1)
    section2 = [
        "Практическая разработка была организована как последовательность взаимосвязанных этапов: подготовка "
        "датасета, настройка цикла обучения, сравнительный запуск моделей, анализ устойчивости, проектирование "
        "демонстрационного интерфейса и формирование экспортируемой отчётности.",
        "При выполнении работы возникла существенная техническая особенность: исходный датасет PKLot поставляется "
        "в виде большого архива PKLot.tar.gz, содержащего изображения и XML-файлы разметки. Случайный доступ к "
        "сжатому архиву приводит к многократной повторной распаковке данных и значительно замедляет предобработку. "
        "Поэтому в рамках практики был реализован последовательный двухпроходный алгоритм чтения архива, "
        "что позволило существенно сократить время формирования выборки.",
        "Выполненная реализация предусматривает не только обучение моделей, но и прикладной режим использования. "
        "Пользователь может загрузить изображение парковочного места, получить предсказанный класс и уверенность "
        "модели, просмотреть накопленную историю проверок, а также скачать Excel-отчёт со сводными метриками. "
        "Тем самым фактически выполненная работа соответствует логике полноценного демонстрационного модуля.",
    ]
    for block in section2:
        add_text(doc, block)

    add_heading(doc, "2.1. Использованные данные", level=2)
    add_text(
        doc,
        "В качестве исходной информационно-эмпирической базы был использован датасет PKLot, включающий более "
        "12 тысяч кадров парковок и порядка 695 тысяч сегментированных парковочных мест [6]. Для текущего этапа "
        "работы была сформирована сбалансированная подвыборка объёмом 10 000 изображений: по 5 000 примеров "
        'классов "empty" и "occupied". Разбиение на train, validation и test выполнялось стратифицированно.',
    )
    add_table(
        doc,
        ["Подмножество", "Класс", "Количество"],
        [[row["split"], row["label"], str(int(row["count"]))] for _, row in split_df.iterrows()],
    )

    figure_number = 1
    figure_number = add_picture(
        doc,
        ASSETS_DIR / "dataset_distribution.png",
        14.0,
        figure_number,
        "Распределение изображений по обучающему, валидационному и тестовому подмножествам.",
    )
    figure_number = add_picture(
        doc,
        ASSETS_DIR / "dataset_samples.png",
        15.0,
        figure_number,
        "Примеры сегментированных изображений парковочных мест после автоматической подготовки выборки.",
    )

    add_heading(doc, "2.2. Методы и инструменты", level=2)
    methods_blocks = [
        "Разработка программной системы выполнена на языке Python 3.11 [7]. Для построения и дообучения моделей "
        "использовалась библиотека PyTorch [8], а готовые архитектуры и предобученные веса загружались из "
        "torchvision [9]. Для расчёта accuracy, precision, recall, F1-score и confusion matrix использовались "
        "средства scikit-learn [10]. Для табличной обработки результатов и экспорта данных в формат Excel "
        "применялась библиотека pandas [11]. Для операций загрузки, вырезания и сохранения изображений "
        "использовалась библиотека Pillow [12].",
        "Сравнение было проведено по пяти различным архитектурам: ResNet18, MobileNetV3 Small, EfficientNet-B0, "
        "DenseNet121 и ConvNeXt-Tiny. Такой выбор обеспечивает сопоставление как компактных, так и более тяжёлых "
        "моделей, что позволяет оценить компромисс между точностью классификации и вычислительными затратами.",
    ]
    for block in methods_blocks:
        add_text(doc, block)

    add_table(doc, ["Параметр", "Значение", "Комментарий"], training_rows)

    add_text(
        doc,
        "Следует отдельно отметить, что обучение выполнялось в ускоренном экспериментальном режиме: по две эпохи "
        "для каждой модели. Это решение было принято как практический компромисс, позволяющий в ограниченное время "
        "получить реальную сравнительную таблицу результатов, протестировать корректность всей инфраструктуры и "
        "подготовить основу для дальнейшего увеличения числа эпох в следующей итерации экспериментов.",
    )

    figure_number = add_picture(
        doc,
        ASSETS_DIR / "training_pipeline.png",
        16.0,
        figure_number,
        "Основные этапы реализованного конвейера: от архива PKLot до демонстрационного модуля и Excel-отчёта.",
    )

    add_heading(doc, "2.3. Ограничения и допущения", level=2)
    limitations = [
        "В работе использовалась подвыборка из 10 000 изображений, а не полный объём PKLot. Это допущение "
        "позволило ускорить подготовку данных и сделать эксперимент воспроизводимым в рамках производственной практики.",
        "Оценка моделей выполнялась на дневных изображениях, представленных в PKLot. Ночные сцены и крайние "
        "условия низкой освещённости в текущем эксперименте не рассматривались.",
        "Форма демонстрационного приложения выбрана в виде локального Streamlit-интерфейса. Этого достаточно "
        "для демонстрации практической работоспособности решения и выполнения требований задания, хотя в будущем "
        "может быть реализован отдельный веб-сервис или настольное приложение.",
    ]
    for block in limitations:
        add_text(doc, block)

    add_page_break(doc)
    add_heading(doc, "3. Описание результатов практической деятельности", level=1)
    add_text(
        doc,
        "На третьем этапе была проведена серия вычислительных экспериментов по сравнению пяти архитектур "
        "нейронных сетей. Результаты тестовой выборки показывают, что все рассмотренные модели обеспечивают "
        "качество классификации на уровне около 96.8-97.0% accuracy, однако различаются по скорости инференса "
        "и размеру файлов моделей.",
    )

    add_table(
        doc,
        ["Модель", "Accuracy", "Precision", "Recall", "F1", "мс/изобр.", "МБ"],
        [
            [
                row["model_name"],
                f"{row['accuracy']:.4f}",
                f"{row['precision']:.4f}",
                f"{row['recall']:.4f}",
                f"{row['f1']:.4f}",
                f"{row['latency_ms_per_image']:.2f}",
                f"{row['model_size_mb']:.2f}",
            ]
            for _, row in summary_df.iterrows()
        ],
    )

    figure_number = add_picture(
        doc,
        ASSETS_DIR / "model_comparison.png",
        15.0,
        figure_number,
        "Сравнение пяти архитектур по F1-score и средней задержке инференса.",
    )

    add_text(
        doc,
        f"Лучший интегральный результат по метрике F1-score показала модель {best['model_name']}. Она достигла "
        f"accuracy {best['accuracy']:.4f}, precision {best['precision']:.4f}, recall {best['recall']:.4f} и "
        f"F1-score {best['f1']:.4f}. При этом среднее время инференса составило {best['latency_ms_per_image']:.2f} "
        f"мс на одно изображение, что делает данную архитектуру практически пригодной для дальнейшего развития в "
        "сторону прикладной системы мониторинга парковочных мест.",
    )

    figure_number = add_picture(
        doc,
        ASSETS_DIR / "best_model_confusion_matrix.png",
        10.5,
        figure_number,
        f"Матрица ошибок лучшей модели {best['model_name']} на тестовой выборке.",
    )

    add_heading(doc, "3.1. Анализ устойчивости к погодным условиям", level=2)
    add_text(
        doc,
        "Одним из обязательных аспектов темы является анализ устойчивости решения к различным условиям наблюдения. "
        "Для этого была отдельно рассчитана точность лучшей модели по погодным категориям sunny, cloudy и rainy. "
        "Расчёт выполнялся на тестовой выборке, что позволило оценить, в каких условиях модель работает стабильнее, "
        "а где качество падает.",
    )
    add_table(
        doc,
        ["Погода", "Образцов", "Accuracy", "Precision", "Recall", "F1"],
        [
            [
                row["weather"],
                str(int(row["samples"])),
                f"{row['accuracy']:.4f}",
                f"{row['precision']:.4f}",
                f"{row['recall']:.4f}",
                f"{row['f1']:.4f}",
            ]
            for _, row in weather_df.iterrows()
        ],
    )
    figure_number = add_picture(
        doc,
        ASSETS_DIR / "weather_comparison.png",
        15.0,
        figure_number,
        "Сравнение качества лучшей модели по погодным условиям.",
    )
    add_text(
        doc,
        "Полученные результаты показывают, что наибольшее число ошибок наблюдается в подмножестве sunny. Для "
        "облачных и дождливых сцен в текущей тестовой подвыборке модель показала идеальные значения метрик, "
        "тогда как солнечные изображения содержали ложные срабатывания и пропуски. Это можно объяснить высоким "
        "контрастом, бликами, резкими тенями и локальными перепадами яркости, которые искажают визуальные признаки "
        "свободного и занятого парковочного места.",
    )

    add_heading(doc, "3.2. Анализ по источникам данных и ракурсам", level=2)
    add_text(
        doc,
        "Помимо погодного анализа был выполнен анализ по источникам данных. В датасете PKLot присутствуют три "
        "источника: pucpr, ufpr04 и ufpr05, отличающиеся ракурсом съёмки и конфигурацией парковки. Разделение "
        "по этим источникам позволяет косвенно оценить устойчивость модели к изменению точки наблюдения.",
    )
    add_table(
        doc,
        ["Источник", "Образцов", "Accuracy", "Precision", "Recall", "F1"],
        [
            [
                row["source"],
                str(int(row["samples"])),
                f"{row['accuracy']:.4f}",
                f"{row['precision']:.4f}",
                f"{row['recall']:.4f}",
                f"{row['f1']:.4f}",
            ]
            for _, row in source_df.iterrows()
        ],
    )
    figure_number = add_picture(
        doc,
        ASSETS_DIR / "source_comparison.png",
        15.0,
        figure_number,
        "Точность лучшей модели по различным парковочным площадкам и ракурсам съёмки.",
    )
    add_text(
        doc,
        "Наиболее сложным для модели оказался источник pucpr, где были сосредоточены почти все ошибки. Источники "
        "ufpr04 и ufpr05 показали существенно более высокое качество. Это означает, что устойчивость решения "
        "зависит не только от погоды, но и от визуальной специфики площадки: перспективного искажения, структуры "
        "разметки, окружающего фона и распределения теней.",
    )

    add_heading(doc, "3.3. Демонстрационный модуль", level=2)
    demo_blocks = [
        "Для выполнения прикладной части задания был реализован демонстрационный модуль на Streamlit. Интерфейс "
        "позволяет загрузить изображение парковочного места, выбрать модель, запустить предсказание, просмотреть "
        "вероятности классов, историю выполненных проверок и краткую статистику по уже полученным результатам.",
        "Важным элементом реализации является сохранение истории запусков. Каждое предсказание записывается в "
        "файл prediction_history.jsonl, где хранятся временная метка, имя модели, предсказанный класс и значения "
        "вероятностей. Тем самым выполнено требование задания о сохранении результатов работы системы.",
        "Кроме этого, сформирован Excel-отчёт parking_occupancy_results.xlsx, содержащий сводные таблицы по "
        "моделям, погоде, распределению выборки, конфигурации обучения и примерам прогнозов. Это закрывает "
        "требование задания о генерации краткого отчёта в PDF или Excel.",
    ]
    for block in demo_blocks:
        add_text(doc, block)
    figure_number = add_picture(
        doc,
        ASSETS_DIR / "streamlit_demo.png",
        16.0,
        figure_number,
        "Интерфейс демонстрационного Streamlit-модуля для определения занятости парковочного места.",
    )
    figure_number = add_picture(
        doc,
        ASSETS_DIR / "prediction_examples.png",
        14.0,
        figure_number,
        "Примеры предсказаний лучшей модели на тестовых изображениях.",
    )

    add_heading(doc, "3.4. Практический результат и направления развития", level=2)
    add_text(
        doc,
        "Итогом практической деятельности стал локальный программный проект, включающий сценарий подготовки "
        "датасета, модуль сравнения пяти моделей, средство генерации графических и табличных отчётов, а также "
        "демонстрационный пользовательский интерфейс. Такой состав проекта позволяет не только показать результат "
        "в рамках отчёта, но и продолжить развитие системы в последующих итерациях.",
    )
    add_text(
        doc,
        "Дальнейшее развитие проекта целесообразно вести по следующим направлениям: увеличение числа эпох "
        "обучения; расширение анализа ошибок на более крупной тестовой выборке; исследование устойчивости при "
        "изменении масштаба и геометрии парковочного места; добавление отчётности в PDF; размещение итогового "
        "проекта в публичном репозитории GitHub и автоматизация воспроизводимого запуска.",
    )

    add_page_break(doc)
    add_heading(doc, "Заключение", level=1)
    conclusion_blocks = [
        "В ходе производственной практики была решена прикладная задача определения занятости парковочного места "
        "по изображению. Выполнена подготовка данных на основе PKLot, разработан ускоренный конвейер обработки "
        "архива, проведено сравнение пяти архитектур нейронных сетей и выбрана лучшая модель по совокупности "
        "метрик качества и скорости.",
        "Фактически выполненная работа показала, что даже в ускоренном экспериментальном режиме возможно получить "
        "устойчивый и воспроизводимый базовый результат. Лучшей архитектурой оказалась EfficientNet-B0, которая "
        "обеспечила наибольший F1-score при умеренном размере модели.",
        "Дополнительно была реализована прикладная часть проекта: демонстрационный Streamlit-интерфейс, ведение "
        "истории предсказаний и генерация сводного Excel-отчёта. Это делает результат практики не только "
        "исследовательским, но и инженерно завершённым на уровне учебного проекта.",
    ]
    for block in conclusion_blocks:
        add_text(doc, block)

    add_page_break(doc)
    add_heading(doc, "Список использованных источников", level=1)
    sources = [
        "Marek M. Image-Based Parking Space Occupancy Classification: Dataset and Baseline [Электронный ресурс]. 2021. URL: https://arxiv.org/abs/2107.12207 (дата обращения: 02.07.2026).",
        "Martynova A., Kuznetsov M., Porvatov V., Tishin V., Kuznetsov A., Semenova N., Kuznetsova K. Revising deep learning methods in parking lot occupancy detection [Электронный ресурс]. 2023. URL: https://arxiv.org/abs/2306.04288 (дата обращения: 02.07.2026).",
        "Yuldashev Y., Mukhiddinov M., Abdusalomov A.B., Nasimov R., Cho J. Parking Lot Occupancy Detection with Improved MobileNetV3 // Sensors. 2023. Vol. 23. No. 17. Art. 7642. URL: https://www.mdpi.com/1424-8220/23/17/7642 (дата обращения: 02.07.2026).",
        "Grbic R., Koch B. Automatic Vision-Based Parking Slot Detection and Occupancy Classification [Электронный ресурс]. 2023. URL: https://arxiv.org/abs/2308.08192 (дата обращения: 02.07.2026).",
        "da Luz G.P.C.P., Sato G.M., Gonzalez L.F.G., Borin J.F. Smart Parking with Pixel-Wise ROI Selection for Vehicle Detection Using YOLOv8, YOLOv9, YOLOv10, and YOLOv11 [Электронный ресурс]. 2024. URL: https://arxiv.org/abs/2412.01983 (дата обращения: 02.07.2026).",
        "Parking Lot Database [Электронный ресурс]. Vision, Robotics and Imaging Laboratory, UFPR. URL: https://web.inf.ufpr.br/vri/databases/parking-lot-database/ (дата обращения: 02.07.2026).",
        "Python 3 Documentation [Электронный ресурс]. URL: https://docs.python.org/ (дата обращения: 02.07.2026).",
        "PyTorch Documentation [Электронный ресурс]. URL: https://docs.pytorch.org/docs/stable/index.html (дата обращения: 02.07.2026).",
        "Torchvision Models and Pre-trained Weights [Электронный ресурс]. URL: https://docs.pytorch.org/vision/main/models.html (дата обращения: 02.07.2026).",
        "Scikit-learn. Metrics and scoring: quantifying the quality of predictions [Электронный ресурс]. URL: https://scikit-learn.org/stable/modules/model_evaluation.html (дата обращения: 02.07.2026).",
        "pandas Documentation [Электронный ресурс]. URL: https://pandas.pydata.org/docs/ (дата обращения: 02.07.2026).",
        "Pillow Documentation [Электронный ресурс]. URL: https://pillow.readthedocs.io/ (дата обращения: 02.07.2026).",
    ]
    for index, source in enumerate(sources, start=1):
        add_text(doc, f"{index}. {source}", first_line_indent=False)

    add_page_break(doc)
    add_heading(doc, "Приложения", level=1)
    appendix_blocks = [
        f"Приложение А. Ссылка на репозиторий проекта: {repo_url()}",
        f"Приложение Б. Локальный путь к проекту: {ROOT}",
        f"Приложение В. Итоговый Excel-отчёт: {WORKBOOK_PATH}",
        "Приложение Г. Основные файлы проекта: prepare_pklot.py, benchmark.py, predict.py, export_results_bundle.py, app.py, src/parking_occupancy/models.py, src/parking_occupancy/trainer.py.",
        "Приложение Д. Иллюстративные материалы: dataset_distribution.png, dataset_samples.png, training_pipeline.png, model_comparison.png, weather_comparison.png, source_comparison.png, best_model_confusion_matrix.png, prediction_examples.png, streamlit_demo.png.",
    ]
    for block in appendix_blocks:
        add_text(doc, block, first_line_indent=False)

    doc.save(REPORT_PATH)


if __name__ == "__main__":
    build_report()
